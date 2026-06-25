import base64
import re
from io import BytesIO
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_cabinet_header(doc: Document, cabinet: dict | None):
    """Insère logo + infos cabinet en haut du document Word."""
    if not cabinet:
        return

    has_logo = cabinet.get("logo") and cabinet.get("logoMime")
    has_info = any(cabinet.get(k) for k in ("nomCabinet", "adresse", "telephone", "email", "siteWeb"))

    if not has_logo and not has_info:
        return

    if has_logo:
        p = doc.add_paragraph()
        run = p.add_run()
        try:
            img_bytes = base64.b64decode(cabinet["logo"])
            run.add_picture(BytesIO(img_bytes), height=Inches(0.6))
        except Exception:
            pass
        p.paragraph_format.space_after = Pt(2)

    if has_info:
        info_parts = []
        if cabinet.get("nomCabinet"):
            info_parts.append(("bold", cabinet["nomCabinet"]))
        for key in ("adresse", "telephone", "email", "siteWeb"):
            if cabinet.get(key):
                info_parts.append(("normal", cabinet[key]))

        p = doc.add_paragraph()
        for i, (style_type, text) in enumerate(info_parts):
            if i > 0:
                p.add_run("  ·  ").font.size = Pt(8)
            run = p.add_run(text)
            run.font.size = Pt(8)
            run.bold = (style_type == "bold")
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Ligne de séparation
    sep = doc.add_paragraph()
    sep.paragraph_format.space_after = Pt(8)
    pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A376C')
    pBdr.append(bottom)
    pPr.append(pBdr)


def markdown_to_docx(title: str, markdown_content: str, cabinet: dict | None = None) -> BytesIO:
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # Entête cabinet
    _add_cabinet_header(doc, cabinet)

    # Page de garde
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.runs[0].font.color.rgb = RGBColor(0x1A, 0x37, 0x6C)
    doc.add_paragraph()

    # Conversion Markdown → DOCX
    lines = markdown_content.split("\n")
    for line in lines:
        line = line.rstrip()

        if line.startswith("### "):
            p = doc.add_heading(line[4:], level=3)
            _style_heading(p, 11, RGBColor(0x2E, 0x6D, 0xA0))

        elif line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
            _style_heading(p, 13, RGBColor(0x1A, 0x37, 0x6C))

        elif line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
            _style_heading(p, 14, RGBColor(0x1A, 0x37, 0x6C))

        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
            p.runs[0].font.size = Pt(10) if p.runs else None

        elif line.startswith("  - ") or line.startswith("  * "):
            doc.add_paragraph(line[4:], style="List Bullet 2")

        elif re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s", "", line)
            doc.add_paragraph(text, style="List Number")

        elif line == "" or line == "---":
            if line == "---":
                doc.add_paragraph().add_run().add_break()
            else:
                doc.add_paragraph()

        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
            run.font.size = Pt(10)

        else:
            p = doc.add_paragraph()
            _parse_inline(p, line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _style_heading(paragraph, size: int, color: RGBColor):
    for run in paragraph.runs:
        run.font.size = Pt(size)
        run.font.color.rgb = color


def _parse_inline(paragraph, text: str):
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
        run.font.size = Pt(10)
